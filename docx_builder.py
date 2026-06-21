"""Renders any PlayDoc sheet into a clearly-themed, submission-safe .docx.

Every document keeps the same Becca banner at the top, but ROTATES through a
set of design themes (different accent colours, page border, and heading style)
so each assignment looks like its own concept. Body stays Times New Roman 12pt
double-spaced, which is submission-safe.
"""
import os
import io
import random
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

STUDENT_NAME = os.environ.get("STUDENT_NAME", "Student Name")
STUDENT_ID = os.environ.get("STUDENT_ID", "Student ID")
BANNER_PATH = os.path.join(os.path.dirname(__file__), "assets", "banner.png")

GREY = RGBColor(0x70, 0x70, 0x70)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _rgb(hex6):
    return RGBColor(int(hex6[0:2], 16), int(hex6[2:4], 16), int(hex6[4:6], 16))


# Each theme keeps a SIMILAR banner (the same banner.png on top) but changes the
# accent colours and the heading look, so every sheet feels like a new concept.
THEMES = [
    {  # Blossom — the original pink + sky blue
        "name": "Blossom", "title": "D6336C", "heading": "1C7ED6",
        "border": "7FD4FF", "divider": "F2A9C4", "banner_fill": "FFE3EE",
        "heading_style": "plain", "accent": "FBD6E4", "page_bg": "FFF6FA",
    },
    {  # Meadow — soft green + blue, heading with a coloured side bar
        "name": "Meadow", "title": "2B8A3E", "heading": "1C7ED6",
        "border": "A9E3B5", "divider": "B2E2BE", "banner_fill": "E6F7EA",
        "heading_style": "bar", "accent": "2B8A3E", "page_bg": "F4FBF6",
    },
    {  # Lavender Sky — plum + blue, heading on a light shaded strip
        "name": "Lavender Sky", "title": "7048E8", "heading": "1C7ED6",
        "border": "C0B6F2", "divider": "D0C8F5", "banner_fill": "F0ECFB",
        "heading_style": "shaded", "accent": "EEE9FB", "page_bg": "F7F4FD",
    },
    {  # Sunrise — coral + teal, heading with an underline
        "name": "Sunrise", "title": "E8590C", "heading": "1098AD",
        "border": "FFC9A0", "divider": "FFD8B0", "banner_fill": "FFF0E6",
        "heading_style": "underline", "accent": "E8590C", "page_bg": "FFF7F0",
    },
]

# Start somewhere random so consecutive runs don't always begin with Blossom,
# then step through one by one so two docs in a row never look the same.
_state = {"i": random.randrange(len(THEMES))}


def _next_theme():
    t = THEMES[_state["i"] % len(THEMES)]
    _state["i"] += 1
    return t


def _para_border(p, edge, color, size="14", space="4"):
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    b = OxmlElement(f"w:{edge}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), size)
    b.set(qn("w:space"), space)
    b.set(qn("w:color"), color)
    pbdr.append(b)


def _para_shade(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _divider(doc, color="F2A9C4", size="10"):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    _para_border(p, "bottom", color, size=size, space="1")
    return p


def _page_background(doc, color):
    """Fill the whole page with a soft colour (shows in Word and Pages)."""
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), color)
    doc.element.insert(0, bg)  # must be the first child of <w:document>
    try:
        settings = doc.settings.element
        if settings.find(qn("w:displayBackgroundShape")) is None:
            settings.append(OxmlElement("w:displayBackgroundShape"))
    except Exception:
        pass


def _page_border(doc, color="7FD4FF"):
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f"w:{edge}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "18")
            b.set(qn("w:space"), "24")
            b.set(qn("w:color"), color)
            pgBorders.append(b)
        sectPr.append(pgBorders)


def _footer_student_id(doc, student_id):
    """Show 'Student id: ...' on the bottom-left of every page."""
    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fr = fp.add_run(f"Student id: {student_id}")
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(9)
    fr.font.color.rgb = GREY


def _cover_page(doc, lines):
    """A title page (page 1 only): the given lines centered, then a page break.
    Uses the section's vertical-centre so it sits in the middle like a real cover."""
    for ln in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(8)
        if str(ln).strip():
            r = p.add_run(str(ln).strip())
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.font.color.rgb = BLACK
    # vertically centre this (cover) section
    sectPr = doc.sections[0]._sectPr
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), "center")
    sectPr.append(vAlign)
    # start a fresh, top-aligned page for the body
    doc.add_section(WD_SECTION.NEW_PAGE)


def _plain_heading(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = h.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = BLACK


def _banner(doc, title, subtitle, theme):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), theme["banner_fill"])
    cell._tc.get_or_add_tcPr().append(shd)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(20)
    r.font.color.rgb = _rgb(theme["title"])

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(subtitle)
    r2.italic = True
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(11)
    r2.font.color.rgb = GREY


def _heading(doc, text, theme):
    """A section heading drawn as a soft colour card, plus a per-theme accent."""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # every heading sits on a light band so it reads like a designed card
    _para_shade(h, theme["banner_fill"])
    h.paragraph_format.left_indent = Pt(6)
    style = theme["heading_style"]
    if style == "bar":
        _para_border(h, "left", theme["accent"], size="24", space="6")
    elif style == "underline":
        _para_border(h, "bottom", theme["accent"], size="8", space="2")
    r = h.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = _rgb(theme["heading"])
    return h


def build_docx(title, subtitle, sections, path, doc_type=False, photo_bytes=None,
               theme=None, decor=None, style="plain", cover_lines=None, student_id=None):
    """Default style is 'plain': white page, no banner, plain black text, like a
    normal hand-in. style='pretty' brings back the Becca banner, soft colour, and
    decorations (only when the user asks for it).

    cover_lines: list of strings for a title page (page 1 only).
    student_id: shown bottom-left on every page.
    """
    pretty = (style == "pretty")
    if pretty and theme is None:
        theme = _next_theme()

    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    sid = student_id or (STUDENT_ID if STUDENT_ID != "Student ID" else "")
    if sid:
        _footer_student_id(doc, sid)

    if pretty:
        _page_background(doc, theme.get("page_bg", "FFF6FA"))

    # Title page (first page only), then a fresh page for the body.
    if cover_lines:
        _cover_page(doc, [ln for ln in cover_lines])

    # ---- body ----
    if pretty:
        if os.path.isfile(BANNER_PATH):
            bp = doc.add_paragraph()
            bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bp.paragraph_format.space_after = Pt(4)
            try:
                bp.add_run().add_picture(BANNER_PATH, width=Inches(6.3))
            except Exception:
                _banner(doc, title, subtitle, theme)
        if title:
            tt = doc.add_paragraph()
            tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            ttr = tt.add_run(title)
            ttr.bold = True
            ttr.font.name = "Times New Roman"
            ttr.font.size = Pt(16)
            ttr.font.color.rgb = _rgb(theme["title"])
        if subtitle:
            sp = doc.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            spr = sp.add_run(subtitle)
            spr.italic = True
            spr.font.name = "Times New Roman"
            spr.font.size = Pt(11)
            spr.font.color.rgb = GREY
        if decor:
            dp = doc.add_paragraph()
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            dr = dp.add_run(decor)
            dr.font.size = Pt(15)
        _divider(doc, color=theme["border"], size="12")
    else:
        # plain: a simple black title, no banner, no colour
        if title:
            tt = doc.add_paragraph()
            tt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            tt.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            tt.paragraph_format.space_after = Pt(6)
            ttr = tt.add_run(title)
            ttr.bold = True
            ttr.font.name = "Times New Roman"
            ttr.font.size = Pt(14)
            ttr.font.color.rgb = BLACK

    for heading, body in sections:
        if heading:
            if pretty:
                _heading(doc, heading, theme)
            else:
                _plain_heading(doc, heading)
        for line in str(body).split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

    if photo_bytes:
        cap = doc.add_paragraph()
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        cr = cap.add_run("Photo")
        cr.bold = True
        cr.font.name = "Times New Roman"
        cr.font.size = Pt(12)
        cr.font.color.rgb = _rgb(theme["heading"]) if pretty else BLACK
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            pic.add_run().add_picture(io.BytesIO(photo_bytes), width=Inches(4.5))
        except Exception:
            pass

    if pretty:
        _page_border(doc, color=theme["border"])

    doc.save(path)
    return path
